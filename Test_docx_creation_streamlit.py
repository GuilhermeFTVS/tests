import streamlit as st
import docx
from docx.shared import Inches
from io import BytesIO


# Function to generate the report
def generate_report(name, age, city):
    # Create a new Word document
    document = docx.Document()

    # Add a heading to the document
    document.add_heading("User Information Report", 0)

    # Add some text to the document
    document.add_paragraph("Name: " + name)
    document.add_paragraph("Age: " + age)
    document.add_paragraph("City: " + city)

    # Save the document to a BytesIO object
    output = BytesIO()
    document.save(output)
    output.seek(0)

    return output


# Streamlit app
def main():
    st.title("Generate Word Report")
    st.write("Enter your information below to generate a report.")

    # User inputs
    name = st.text_input("Name")
    age = st.text_input("Age")
    city = st.text_input("City")

    # Generate the report
    if st.button("Generate Report"):
        report = generate_report(name, age, city)

        # Download the report
        st.download_button(
            label="Download Report",
            data=report,
            file_name="user_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )


if __name__ == "__main__":
    main()
